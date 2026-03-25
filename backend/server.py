import os
import json
import uuid
from datetime import datetime, timedelta, date
from typing import Optional, List

import psycopg2
import psycopg2.extras
from psycopg2 import pool as pg_pool
import bcrypt
import jwt
from fastapi import FastAPI, HTTPException, Depends, Header, Response, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
from io import BytesIO

load_dotenv()

app = FastAPI(title="Simondu Web API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

DATABASE_URL = os.environ.get("DATABASE_URL")
JWT_SECRET = os.environ.get("JWT_SECRET", "simondu-web-secret-key-2024-polda-jabar")
JWT_ALGORITHM = os.environ.get("JWT_ALGORITHM", "HS256")
JWT_EXPIRATION_MINUTES = int(os.environ.get("JWT_EXPIRATION_MINUTES", "480"))

# ==================== DB HELPERS ====================
db_pool = None

def init_pool():
    global db_pool
    if db_pool is None:
        db_pool = pg_pool.ThreadedConnectionPool(2, 15, DATABASE_URL)
    return db_pool

def get_db():
    p = init_pool()
    conn = p.getconn()
    conn.autocommit = False
    return conn

def release_db(conn):
    try:
        p = init_pool()
        p.putconn(conn)
    except Exception:
        try: conn.close()
        except: pass

def serialize_row(row, columns):
    if row is None:
        return None
    result = {}
    for i, col in enumerate(columns):
        val = row[i]
        if isinstance(val, (datetime, date)):
            result[col] = val.isoformat()
        elif isinstance(val, uuid.UUID):
            result[col] = str(val)
        else:
            result[col] = val
    return result

def query_all(sql, params=None):
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute(sql, params or ())
        columns = [desc[0] for desc in cur.description]
        rows = cur.fetchall()
        cur.close()
        return [serialize_row(r, columns) for r in rows]
    finally:
        release_db(conn)

def query_one(sql, params=None):
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute(sql, params or ())
        if cur.description is None:
            cur.close()
            return None
        columns = [desc[0] for desc in cur.description]
        row = cur.fetchone()
        cur.close()
        return serialize_row(row, columns) if row else None
    finally:
        release_db(conn)

def execute(sql, params=None):
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute(sql, params or ())
        conn.commit()
        if cur.description:
            columns = [desc[0] for desc in cur.description]
            row = cur.fetchone()
            cur.close()
            return serialize_row(row, columns) if row else None
        cur.close()
        return None
    except Exception as e:
        conn.rollback()
        raise e
    finally:
        release_db(conn)

# ==================== AUDIT LOG ====================
def log_audit(user, action, entity_type=None, entity_id=None, details=None):
    try:
        execute(
            "INSERT INTO audit_log (user_id, user_name, user_role, action, entity_type, entity_id, details) VALUES (%s::uuid, %s, %s, %s, %s, %s, %s)",
            (user.get("sub"), user.get("name"), user.get("role"), action, entity_type, str(entity_id) if entity_id else None, details)
        )
    except Exception:
        pass

# ==================== AUTH ====================
def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(password: str, hashed: str) -> bool:
    return bcrypt.checkpw(password.encode(), hashed.encode())

def create_token(user_id, role, email, name, unit_name=None):
    payload = {
        "sub": user_id, "role": role, "email": email, "name": name,
        "unit_name": unit_name,
        "exp": datetime.utcnow() + timedelta(minutes=JWT_EXPIRATION_MINUTES),
        "iat": datetime.utcnow()
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def get_current_user(authorization: Optional[str] = Header(None)):
    if not authorization:
        raise HTTPException(status_code=401, detail="Token tidak ditemukan")
    try:
        token = authorization.replace("Bearer ", "")
        return jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token sudah expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Token tidak valid")

def require_role(*roles):
    def checker(user=Depends(get_current_user)):
        if user["role"] not in roles:
            raise HTTPException(status_code=403, detail="Akses ditolak")
        return user
    return checker

# ==================== PYDANTIC MODELS ====================
class LoginRequest(BaseModel):
    email: str
    password: str

class DumasCreate(BaseModel):
    no_dumas: Optional[str] = None
    tgl_dumas: Optional[str] = None
    pelapor: Optional[str] = None
    terlapor: Optional[str] = None
    satker: Optional[str] = None
    wujud_perbuatan: Optional[str] = None
    jenis_dumas: Optional[str] = None
    keterangan: Optional[str] = None
    disposisi_kabid: Optional[str] = None
    disposisi_kasubbid: Optional[str] = None
    unit_id: Optional[str] = None
    status: Optional[str] = "dalam_proses"
    parent_dumas_id: Optional[str] = None

class DumasUpdate(BaseModel):
    no_dumas: Optional[str] = None
    tgl_dumas: Optional[str] = None
    pelapor: Optional[str] = None
    terlapor: Optional[str] = None
    satker: Optional[str] = None
    wujud_perbuatan: Optional[str] = None
    jenis_dumas: Optional[str] = None
    keterangan: Optional[str] = None
    disposisi_kabid: Optional[str] = None
    disposisi_kasubbid: Optional[str] = None
    unit_id: Optional[str] = None
    status: Optional[str] = None
    parent_dumas_id: Optional[str] = None

class TindakLanjutUpsert(BaseModel):
    dumas_id: str
    no_berkas: Optional[str] = None
    tgl_uuk: Optional[str] = None
    no_uuk: Optional[str] = None
    tgl_sprin_lidik: Optional[str] = None
    no_sprin: Optional[str] = None
    tgl_gelar: Optional[str] = None
    tgl_lhp: Optional[str] = None
    no_lhp: Optional[str] = None
    hasil_lidik: Optional[str] = None
    tgl_pelimpahan: Optional[str] = None
    no_pelimpahan: Optional[str] = None
    tgl_henti_lidik: Optional[str] = None
    no_henti_lidik: Optional[str] = None
    tgl_nodin: Optional[str] = None
    no_nodin: Optional[str] = None
    tgl_sp2hp2: Optional[str] = None
    no_sp2hp2: Optional[str] = None
    tgl_ke_ankum: Optional[str] = None
    tgl_ke_mabes: Optional[str] = None
    no_mabes: Optional[str] = None
    tgl_st_arahan: Optional[str] = None
    no_arahan: Optional[str] = None
    catatan: Optional[str] = None

class ApprovalAction(BaseModel):
    action: str
    catatan_revisi: Optional[str] = None

class PenyelesaianCreate(BaseModel):
    dumas_id: str
    jenis: str
    tanggal: Optional[str] = None
    nomor: Optional[str] = None

class TimLidikCreate(BaseModel):
    dumas_id: str
    anggota_id: str
    no_hp_penyelidik: Optional[str] = None

class SettingUpdate(BaseModel):
    value: list | dict

class BulkDumasAction(BaseModel):
    dumas_ids: List[str]
    action: str  # "merge" or "delete"
    parent_dumas_id: Optional[str] = None

class UserCreate(BaseModel):
    email: str
    password: str
    name: str
    role: str
    unit_name: Optional[str] = None

class UserUpdate(BaseModel):
    name: Optional[str] = None
    role: Optional[str] = None
    unit_name: Optional[str] = None
    password: Optional[str] = None

# ==================== AUTO NUMBERING ====================
ROMAN_MONTHS = {1:"I",2:"II",3:"III",4:"IV",5:"V",6:"VI",7:"VII",8:"VIII",9:"IX",10:"X",11:"XI",12:"XII"}

def auto_number(format_template, nomor, tanggal=None, unit=None):
    if tanggal is None:
        tanggal = date.today()
    if isinstance(tanggal, str):
        tanggal = datetime.strptime(tanggal, "%Y-%m-%d").date()
    result = format_template.replace("{nomor}", str(nomor).zfill(3))
    result = result.replace("{bulan_romawi}", ROMAN_MONTHS.get(tanggal.month, "I"))
    result = result.replace("{tahun}", str(tanggal.year))
    result = result.replace("{unit}", unit or "PAMINAL")
    return result

# ==================== ROUTES ====================

@app.get("/api/health")
def health():
    return {"status": "ok", "service": "Simondu Web API v2"}

# --- AUTH ---
@app.post("/api/auth/login")
def login(req: LoginRequest):
    user = query_one("SELECT * FROM users WHERE email = %s AND deleted_at IS NULL", (req.email,))
    if not user or not verify_password(req.password, user["password_hash"]):
        raise HTTPException(status_code=401, detail="Email atau password salah")
    token = create_token(user["id"], user["role"], user["email"], user["name"], user.get("unit_name"))
    return {"token": token, "user": {"id": user["id"], "email": user["email"], "name": user["name"], "role": user["role"], "unit_name": user.get("unit_name")}}

@app.get("/api/auth/me")
def get_me(user=Depends(get_current_user)):
    db_user = query_one("SELECT id, email, name, role, unit_name FROM users WHERE id = %s::uuid", (user["sub"],))
    if not db_user:
        raise HTTPException(status_code=404, detail="User tidak ditemukan")
    return db_user

# --- USERS ---
@app.get("/api/users")
def list_users(user=Depends(get_current_user)):
    return query_all("SELECT id, email, name, role, unit_name, created_at, deleted_at FROM users ORDER BY name")

@app.get("/api/users/units")
def list_unit_users(user=Depends(get_current_user)):
    return query_all("SELECT id, email, name, role, unit_name FROM users WHERE role = 'unit' AND deleted_at IS NULL ORDER BY name")

@app.post("/api/users")
def create_user(req: UserCreate, user=Depends(require_role("superadmin"))):
    pw_hash = hash_password(req.password)
    result = execute(
        "INSERT INTO users (email, password_hash, name, role, unit_name) VALUES (%s, %s, %s, %s, %s) RETURNING id, email, name, role, unit_name",
        (req.email, pw_hash, req.name, req.role, req.unit_name)
    )
    log_audit(user, f"Menambahkan user {req.name} ({req.role})", "users", result.get("id") if result else None)
    return result

@app.put("/api/users/{user_id}")
def update_user(user_id: str, req: UserUpdate, user=Depends(require_role("superadmin"))):
    updates, params = [], []
    if req.name: updates.append("name = %s"); params.append(req.name)
    if req.role: updates.append("role = %s"); params.append(req.role)
    if req.unit_name is not None: updates.append("unit_name = %s"); params.append(req.unit_name)
    if req.password: updates.append("password_hash = %s"); params.append(hash_password(req.password))
    if not updates:
        raise HTTPException(status_code=400, detail="Tidak ada data")
    updates.append("updated_at = NOW()")
    params.append(user_id)
    result = execute(f"UPDATE users SET {', '.join(updates)} WHERE id = %s::uuid RETURNING id, email, name, role, unit_name", tuple(params))
    log_audit(user, f"Mengubah data user {result.get('name') if result else user_id}", "users", user_id)
    return result

@app.delete("/api/users/{user_id}")
def soft_delete_user(user_id: str, user=Depends(require_role("superadmin"))):
    result = execute("UPDATE users SET deleted_at = NOW() WHERE id = %s::uuid RETURNING name", (user_id,))
    if result:
        log_audit(user, f"Menonaktifkan user {result.get('name')}", "users", user_id)
    return {"message": "User berhasil dinonaktifkan"}

@app.post("/api/users/{user_id}/restore")
def restore_user(user_id: str, user=Depends(require_role("superadmin"))):
    execute("UPDATE users SET deleted_at = NULL WHERE id = %s::uuid", (user_id,))
    return {"message": "User berhasil diaktifkan kembali"}

# --- SETTINGS ---
@app.get("/api/settings")
def list_settings(user=Depends(get_current_user)):
    return query_all("SELECT * FROM settings ORDER BY type")

@app.get("/api/settings/{setting_type}")
def get_setting(setting_type: str, user=Depends(get_current_user)):
    setting = query_one("SELECT * FROM settings WHERE type = %s", (setting_type,))
    if not setting:
        raise HTTPException(status_code=404, detail="Setting tidak ditemukan")
    return setting

@app.put("/api/settings/{setting_type}")
def update_setting(setting_type: str, req: SettingUpdate, user=Depends(require_role("superadmin", "admin"))):
    result = execute("UPDATE settings SET value = %s::jsonb, updated_at = NOW() WHERE type = %s RETURNING *", (json.dumps(req.value), setting_type))
    if not result:
        # Create if not exists
        result = execute("INSERT INTO settings (type, value) VALUES (%s, %s::jsonb) RETURNING *", (setting_type, json.dumps(req.value)))
    log_audit(user, f"Mengubah pengaturan {setting_type}", "settings")
    return result

# --- DUMAS ---
@app.get("/api/dumas")
def list_dumas(status: Optional[str] = None, unit_id: Optional[str] = None, search: Optional[str] = None,
               page: int = 1, limit: int = 50, include_deleted: bool = False, user=Depends(get_current_user)):
    sql = "SELECT d.*, u.name as unit_name FROM dumas d LEFT JOIN users u ON d.unit_id = u.id"
    conditions, params = [], []
    if not include_deleted:
        conditions.append("d.deleted_at IS NULL")
    if user["role"] == "unit":
        conditions.append("d.unit_id = %s::uuid")
        params.append(user["sub"])
    elif unit_id:
        conditions.append("d.unit_id = %s::uuid")
        params.append(unit_id)
    if status:
        conditions.append("d.status = %s")
        params.append(status)
    if search:
        conditions.append("(d.no_dumas ILIKE %s OR d.pelapor ILIKE %s OR d.terlapor ILIKE %s OR d.satker ILIKE %s OR d.keterangan ILIKE %s)")
        s = f"%{search}%"
        params.extend([s, s, s, s, s])
    if conditions:
        sql += " WHERE " + " AND ".join(conditions)
    sql += " ORDER BY d.created_at DESC"
    offset = (page - 1) * limit
    sql += f" LIMIT {limit} OFFSET {offset}"
    return query_all(sql, tuple(params))

@app.get("/api/dumas/count")
def count_dumas(status: Optional[str] = None, search: Optional[str] = None, user=Depends(get_current_user)):
    sql = "SELECT COUNT(*) as total FROM dumas d"
    conditions, params = ["d.deleted_at IS NULL"], []
    if user["role"] == "unit":
        conditions.append("d.unit_id = %s::uuid")
        params.append(user["sub"])
    if status:
        conditions.append("d.status = %s")
        params.append(status)
    if search:
        conditions.append("(d.no_dumas ILIKE %s OR d.pelapor ILIKE %s OR d.terlapor ILIKE %s OR d.satker ILIKE %s)")
        s = f"%{search}%"
        params.extend([s, s, s, s])
    sql += " WHERE " + " AND ".join(conditions)
    result = query_one(sql, tuple(params))
    return {"total": result["total"] if result else 0}

@app.post("/api/dumas")
def create_dumas(req: DumasCreate, user=Depends(require_role("admin", "superadmin"))):
    result = execute(
        """INSERT INTO dumas (no_dumas, tgl_dumas, pelapor, terlapor, satker, wujud_perbuatan,
           jenis_dumas, keterangan, disposisi_kabid, disposisi_kasubbid, unit_id, status, parent_dumas_id)
           VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING *""",
        (req.no_dumas, req.tgl_dumas, req.pelapor, req.terlapor, req.satker,
         req.wujud_perbuatan, req.jenis_dumas, req.keterangan,
         req.disposisi_kabid, req.disposisi_kasubbid,
         req.unit_id if req.unit_id else None, req.status or "dalam_proses",
         req.parent_dumas_id if req.parent_dumas_id else None)
    )
    log_audit(user, f"Menambahkan Dumas {req.no_dumas or 'baru'}", "dumas", result.get("id") if result else None)
    return Response(content=json.dumps(result, default=str), status_code=201, media_type="application/json")

@app.get("/api/dumas/{dumas_id}")
def get_dumas(dumas_id: str, user=Depends(get_current_user)):
    dumas = query_one(
        "SELECT d.*, u.name as unit_name FROM dumas d LEFT JOIN users u ON d.unit_id = u.id WHERE d.id = %s::uuid AND d.deleted_at IS NULL",
        (dumas_id,))
    if not dumas:
        raise HTTPException(status_code=404, detail="Dumas tidak ditemukan")
    return dumas

@app.put("/api/dumas/{dumas_id}")
def update_dumas(dumas_id: str, req: DumasUpdate, user=Depends(get_current_user)):
    updates, params = [], []
    data = req.dict(exclude_unset=True, exclude_none=True)
    for key, value in data.items():
        if key in ("unit_id", "parent_dumas_id"):
            updates.append(f"{key} = %s::uuid")
        else:
            updates.append(f"{key} = %s")
        params.append(value)
    if not updates:
        raise HTTPException(status_code=400, detail="Tidak ada data")
    updates.append("updated_at = NOW()")
    params.append(dumas_id)
    result = execute(f"UPDATE dumas SET {', '.join(updates)} WHERE id = %s::uuid AND deleted_at IS NULL RETURNING *", tuple(params))
    if not result:
        raise HTTPException(status_code=404, detail="Dumas tidak ditemukan")
    log_audit(user, f"Mengubah Dumas {result.get('no_dumas', dumas_id)}", "dumas", dumas_id, json.dumps(data))
    return result

@app.delete("/api/dumas/{dumas_id}")
def soft_delete_dumas(dumas_id: str, user=Depends(require_role("admin", "superadmin"))):
    result = execute("UPDATE dumas SET deleted_at = NOW() WHERE id = %s::uuid AND deleted_at IS NULL RETURNING id, no_dumas", (dumas_id,))
    if not result:
        raise HTTPException(status_code=404, detail="Dumas tidak ditemukan")
    log_audit(user, f"Menghapus Dumas {result.get('no_dumas', dumas_id)}", "dumas", dumas_id)
    return {"message": "Dumas berhasil dihapus"}

# --- BULK DUMAS ACTIONS ---
@app.post("/api/dumas/bulk")
def bulk_dumas_action(req: BulkDumasAction, user=Depends(require_role("admin", "superadmin"))):
    conn = get_db()
    try:
        cur = conn.cursor()
        if req.action == "merge" and req.parent_dumas_id:
            for cid in req.dumas_ids:
                if cid != req.parent_dumas_id:
                    cur.execute("UPDATE dumas SET parent_dumas_id = %s::uuid, updated_at = NOW() WHERE id = %s::uuid", (req.parent_dumas_id, cid))
            conn.commit()
            log_audit(user, f"Merge {len(req.dumas_ids)} dumas ke {req.parent_dumas_id}", "dumas")
            return {"message": f"{len(req.dumas_ids)} dumas berhasil di-merge"}
        elif req.action == "delete":
            for did in req.dumas_ids:
                cur.execute("UPDATE dumas SET deleted_at = NOW() WHERE id = %s::uuid", (did,))
            conn.commit()
            log_audit(user, f"Bulk delete {len(req.dumas_ids)} dumas", "dumas")
            return {"message": f"{len(req.dumas_ids)} dumas berhasil dihapus"}
        cur.close()
        raise HTTPException(status_code=400, detail="Action tidak valid")
    except HTTPException:
        raise
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        release_db(conn)

# --- ARCHIVE ---
@app.get("/api/archive/dumas")
def list_archived_dumas(user=Depends(require_role("superadmin"))):
    return query_all("SELECT d.*, u.name as unit_name FROM dumas d LEFT JOIN users u ON d.unit_id = u.id WHERE d.deleted_at IS NOT NULL ORDER BY d.deleted_at DESC")

@app.post("/api/archive/dumas/{dumas_id}/restore")
def restore_dumas(dumas_id: str, user=Depends(require_role("superadmin"))):
    execute("UPDATE dumas SET deleted_at = NULL WHERE id = %s::uuid", (dumas_id,))
    log_audit(user, f"Restore Dumas {dumas_id}", "dumas", dumas_id)
    return {"message": "Dumas berhasil di-restore"}

@app.delete("/api/archive/dumas/{dumas_id}/permanent")
def permanent_delete_dumas(dumas_id: str, user=Depends(require_role("superadmin"))):
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM tim_lidik WHERE dumas_id = %s::uuid", (dumas_id,))
        cur.execute("DELETE FROM penyelesaian WHERE dumas_id = %s::uuid", (dumas_id,))
        cur.execute("DELETE FROM tindak_lanjut WHERE dumas_id = %s::uuid", (dumas_id,))
        cur.execute("UPDATE dumas SET parent_dumas_id = NULL WHERE parent_dumas_id = %s::uuid", (dumas_id,))
        cur.execute("DELETE FROM dumas WHERE id = %s::uuid", (dumas_id,))
        conn.commit()
        cur.close()
        log_audit(user, f"Hapus permanen Dumas {dumas_id}", "dumas", dumas_id)
        return {"message": "Dumas berhasil dihapus permanen"}
    except Exception as e:
        conn.rollback()
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        release_db(conn)

# --- TINDAK LANJUT ---
@app.get("/api/tindak-lanjut")
def list_tindak_lanjut(dumas_id: Optional[str] = None, user=Depends(get_current_user)):
    if dumas_id:
        return query_all("SELECT * FROM tindak_lanjut WHERE dumas_id = %s::uuid AND deleted_at IS NULL ORDER BY created_at DESC", (dumas_id,))
    return query_all("SELECT * FROM tindak_lanjut WHERE deleted_at IS NULL ORDER BY created_at DESC")

@app.post("/api/tindak-lanjut")
def create_or_update_tindak_lanjut(req: TindakLanjutUpsert, user=Depends(get_current_user)):
    # Check if already exists
    existing = query_one("SELECT id, status_verifikasi FROM tindak_lanjut WHERE dumas_id = %s::uuid AND deleted_at IS NULL ORDER BY created_at DESC LIMIT 1", (req.dumas_id,))
    
    fields = req.dict(exclude={'dumas_id'}, exclude_none=True)
    
    if existing:
        if existing["status_verifikasi"] in ["menunggu_verifikasi", "disetujui"]:
            raise HTTPException(status_code=400, detail="Form terkunci - sedang dalam verifikasi atau sudah disetujui")
        updates = [f"{k} = %s" for k in fields.keys()]
        updates.append("updated_at = NOW()")
        params = list(fields.values()) + [existing["id"]]
        result = execute(f"UPDATE tindak_lanjut SET {', '.join(updates)} WHERE id = %s::uuid RETURNING *", tuple(params))
    else:
        cols = ["dumas_id"] + list(fields.keys())
        vals = [req.dumas_id] + list(fields.values())
        placeholders = ["%s::uuid"] + ["%s"] * len(fields)
        result = execute(
            f"INSERT INTO tindak_lanjut ({', '.join(cols)}) VALUES ({', '.join(placeholders)}) RETURNING *",
            tuple(vals)
        )
    
    # Auto-update dumas status based on hasil_lidik
    if req.hasil_lidik:
        status_map = {"terbukti": "terbukti", "tidak_terbukti": "tidak_terbukti"}
        new_status = status_map.get(req.hasil_lidik.lower().replace(" ", "_"))
        if new_status:
            execute("UPDATE dumas SET status = %s, updated_at = NOW() WHERE id = %s::uuid", (new_status, req.dumas_id))
            log_audit(user, f"Status Dumas diubah ke {new_status} berdasarkan hasil lidik", "dumas", req.dumas_id)
    
    log_audit(user, f"Update tindak lanjut Dumas", "tindak_lanjut", result.get("id") if result else None)
    return result

@app.get("/api/tindak-lanjut/{tl_id}")
def get_tindak_lanjut(tl_id: str, user=Depends(get_current_user)):
    return query_one("SELECT * FROM tindak_lanjut WHERE id = %s::uuid", (tl_id,))

# --- APPROVAL ---
@app.post("/api/tindak-lanjut/{tl_id}/submit")
def submit_for_verification(tl_id: str, user=Depends(get_current_user)):
    existing = query_one("SELECT * FROM tindak_lanjut WHERE id = %s::uuid", (tl_id,))
    if not existing:
        raise HTTPException(status_code=404, detail="Tindak lanjut tidak ditemukan")
    if existing["status_verifikasi"] not in ["draft", "revisi"]:
        raise HTTPException(status_code=400, detail="Hanya draft atau revisi yang bisa diajukan")
    result = execute("UPDATE tindak_lanjut SET status_verifikasi = 'menunggu_verifikasi', updated_at = NOW() WHERE id = %s::uuid RETURNING *", (tl_id,))
    dumas = query_one("SELECT no_dumas FROM dumas WHERE id = %s::uuid", (existing["dumas_id"],))
    no_dumas = dumas["no_dumas"] if dumas else "Unknown"
    execute("INSERT INTO notifications (target_role, title, message, action_url) VALUES ('pimpinan', %s, %s, %s)",
        (f"Verifikasi - {no_dumas}", f"Tindak lanjut Dumas {no_dumas} membutuhkan persetujuan", f"/dumas/{existing['dumas_id']}"))
    log_audit(user, f"Mengajukan verifikasi tindak lanjut Dumas {no_dumas}", "tindak_lanjut", tl_id)
    return result

@app.post("/api/tindak-lanjut/{tl_id}/approve")
def approve_tindak_lanjut(tl_id: str, req: ApprovalAction, user=Depends(require_role("pimpinan", "superadmin"))):
    existing = query_one("SELECT * FROM tindak_lanjut WHERE id = %s::uuid", (tl_id,))
    if not existing:
        raise HTTPException(status_code=404, detail="TL tidak ditemukan")
    if existing["status_verifikasi"] != "menunggu_verifikasi":
        raise HTTPException(status_code=400, detail="Status tidak valid")
    dumas = query_one("SELECT no_dumas, unit_id FROM dumas WHERE id = %s::uuid", (existing["dumas_id"],))
    no_dumas = dumas["no_dumas"] if dumas else "Unknown"
    
    if req.action == "setujui":
        result = execute("UPDATE tindak_lanjut SET status_verifikasi = 'disetujui', updated_at = NOW() WHERE id = %s::uuid RETURNING *", (tl_id,))
        if dumas and dumas.get("unit_id"):
            execute("INSERT INTO notifications (target_user_id, title, message, action_url) VALUES (%s::uuid, %s, %s, %s)",
                (dumas["unit_id"], f"Disetujui - {no_dumas}", f"Tindak lanjut Dumas {no_dumas} telah disetujui", f"/dumas/{existing['dumas_id']}"))
        log_audit(user, f"Menyetujui tindak lanjut Dumas {no_dumas}", "tindak_lanjut", tl_id)
    elif req.action == "revisi":
        result = execute("UPDATE tindak_lanjut SET status_verifikasi = 'revisi', catatan_revisi = %s, updated_at = NOW() WHERE id = %s::uuid RETURNING *", (req.catatan_revisi, tl_id))
        if dumas and dumas.get("unit_id"):
            execute("INSERT INTO notifications (target_user_id, title, message, action_url) VALUES (%s::uuid, %s, %s, %s)",
                (dumas["unit_id"], f"Revisi - {no_dumas}", f"Tindak lanjut Dumas {no_dumas} perlu direvisi: {req.catatan_revisi}", f"/dumas/{existing['dumas_id']}"))
        log_audit(user, f"Mengembalikan tindak lanjut Dumas {no_dumas} untuk revisi", "tindak_lanjut", tl_id)
    else:
        raise HTTPException(status_code=400, detail="Action harus 'setujui' atau 'revisi'")
    return result

@app.get("/api/approval/inbox")
def approval_inbox(user=Depends(require_role("pimpinan", "superadmin"))):
    return query_all(
        """SELECT tl.*, d.no_dumas, d.pelapor, d.terlapor, d.satker, d.jenis_dumas, d.keterangan, u.name as unit_name
           FROM tindak_lanjut tl JOIN dumas d ON tl.dumas_id = d.id LEFT JOIN users u ON d.unit_id = u.id
           WHERE tl.status_verifikasi = 'menunggu_verifikasi' AND tl.deleted_at IS NULL ORDER BY tl.updated_at DESC""")

# --- TIM LIDIK ---
@app.post("/api/tim-lidik")
def create_tim_lidik(req: TimLidikCreate, user=Depends(get_current_user)):
    return execute("INSERT INTO tim_lidik (dumas_id, anggota_id, no_hp_penyelidik) VALUES (%s::uuid, %s::uuid, %s) RETURNING *", (req.dumas_id, req.anggota_id, req.no_hp_penyelidik))

@app.get("/api/tim-lidik/{dumas_id}")
def get_tim_lidik(dumas_id: str, user=Depends(get_current_user)):
    return query_all("SELECT tl.*, u.name as anggota_name, u.unit_name FROM tim_lidik tl LEFT JOIN users u ON tl.anggota_id = u.id WHERE tl.dumas_id = %s::uuid", (dumas_id,))

@app.delete("/api/tim-lidik/{tl_id}")
def delete_tim_lidik(tl_id: str, user=Depends(get_current_user)):
    execute("DELETE FROM tim_lidik WHERE id = %s::uuid", (tl_id,))
    return {"message": "OK"}

# --- PENYELESAIAN ---
@app.post("/api/penyelesaian")
def create_penyelesaian(req: PenyelesaianCreate, user=Depends(get_current_user)):
    result = execute("INSERT INTO penyelesaian (dumas_id, jenis, tanggal, nomor) VALUES (%s::uuid, %s, %s, %s) RETURNING *", (req.dumas_id, req.jenis, req.tanggal, req.nomor))
    log_audit(user, f"Menambahkan penyelesaian ({req.jenis}) untuk Dumas", "penyelesaian", result.get("id") if result else None)
    return result

@app.get("/api/penyelesaian/{dumas_id}")
def get_penyelesaian(dumas_id: str, user=Depends(get_current_user)):
    return query_all("SELECT * FROM penyelesaian WHERE dumas_id = %s::uuid ORDER BY created_at", (dumas_id,))

# --- NOTIFICATIONS ---
@app.get("/api/notifications")
def list_notifications(user=Depends(get_current_user)):
    return query_all("SELECT * FROM notifications WHERE (target_user_id = %s::uuid OR target_role = %s) ORDER BY created_at DESC LIMIT 50", (user["sub"], user["role"]))

@app.get("/api/notifications/unread-count")
def unread_count(user=Depends(get_current_user)):
    result = query_one("SELECT COUNT(*) as count FROM notifications WHERE (target_user_id = %s::uuid OR target_role = %s) AND is_read = FALSE", (user["sub"], user["role"]))
    return {"count": result["count"] if result else 0}

@app.put("/api/notifications/{notif_id}/read")
def mark_read(notif_id: str, user=Depends(get_current_user)):
    execute("UPDATE notifications SET is_read = TRUE WHERE id = %s::uuid", (notif_id,))
    return {"message": "OK"}

@app.put("/api/notifications/read-all")
def mark_all_read(user=Depends(get_current_user)):
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("UPDATE notifications SET is_read = TRUE WHERE (target_user_id = %s::uuid OR target_role = %s) AND is_read = FALSE", (user["sub"], user["role"]))
        conn.commit()
        cur.close()
    finally:
        release_db(conn)
    return {"message": "OK"}

# --- DASHBOARD COMBINED ---
@app.get("/api/dashboard/combined")
def dashboard_combined(user=Depends(get_current_user)):
    conn = get_db()
    try:
        cur = conn.cursor()
        
        # Stats
        cur.execute("""SELECT 
            COUNT(*) FILTER (WHERE deleted_at IS NULL) as total,
            COUNT(*) FILTER (WHERE status = 'dalam_proses' AND deleted_at IS NULL) as dalam_proses,
            COUNT(*) FILTER (WHERE status = 'terbukti' AND deleted_at IS NULL) as terbukti,
            COUNT(*) FILTER (WHERE status = 'tidak_terbukti' AND deleted_at IS NULL) as tidak_terbukti
            FROM dumas""")
        cols = [d[0] for d in cur.description]
        stats = serialize_row(cur.fetchone(), cols) or {}
        
        cur.execute("SELECT COUNT(*) FROM tindak_lanjut WHERE status_verifikasi = 'menunggu_verifikasi' AND deleted_at IS NULL")
        tl_count = cur.fetchone()[0]
        
        # By jenis dumas
        cur.execute("""SELECT jenis_dumas as label, COUNT(*) as value FROM dumas 
            WHERE deleted_at IS NULL AND jenis_dumas IS NOT NULL GROUP BY jenis_dumas ORDER BY value DESC LIMIT 10""")
        by_jenis = [serialize_row(r, [d[0] for d in cur.description]) for r in cur.fetchall()]
        
        # By satker
        cur.execute("""SELECT satker as label, COUNT(*) as value FROM dumas 
            WHERE deleted_at IS NULL AND satker IS NOT NULL GROUP BY satker ORDER BY value DESC LIMIT 10""")
        by_satker = [serialize_row(r, [d[0] for d in cur.description]) for r in cur.fetchall()]
        
        # Per unit stats
        cur.execute("""SELECT u.name as unit_name, u.id as unit_id,
            COUNT(d.id) as jumlah_laporan,
            COUNT(d.id) FILTER (WHERE d.status = 'terbukti') as terbukti,
            COUNT(d.id) FILTER (WHERE d.status = 'tidak_terbukti') as tidak_terbukti,
            COUNT(d.id) FILTER (WHERE d.status = 'dalam_proses') as proses,
            CASE WHEN COUNT(d.id) > 0 THEN ROUND(
                (COUNT(d.id) FILTER (WHERE d.status IN ('terbukti','tidak_terbukti'))::numeric / COUNT(d.id) * 100), 1
            ) ELSE 0 END as persentase
            FROM users u LEFT JOIN dumas d ON d.unit_id = u.id AND d.deleted_at IS NULL
            WHERE u.role = 'unit' AND u.deleted_at IS NULL
            GROUP BY u.id, u.name ORDER BY u.name""")
        cols_unit = [d[0] for d in cur.description]
        per_unit = [serialize_row(r, cols_unit) for r in cur.fetchall()]
        
        # Monitoring (latest 10)
        cur.execute("""SELECT d.id, d.no_dumas, d.pelapor, d.terlapor, d.keterangan as perihal,
            d.satker, d.status, d.created_at, d.tgl_dumas,
            u.name as unit_name,
            EXTRACT(DAY FROM NOW() - d.created_at)::int as durasi,
            CASE WHEN EXTRACT(DAY FROM NOW() - d.created_at) > 30 THEN 'merah'
                 WHEN EXTRACT(DAY FROM NOW() - d.created_at) > 14 THEN 'kuning' ELSE 'hijau' END as sla_status
            FROM dumas d LEFT JOIN users u ON d.unit_id = u.id
            WHERE d.deleted_at IS NULL ORDER BY d.created_at DESC LIMIT 100""")
        cols_mon = [d[0] for d in cur.description]
        monitoring = [serialize_row(r, cols_mon) for r in cur.fetchall()]
        
        cur.close()
        return {
            "stats": {**stats, "menunggu_verifikasi": tl_count or 0},
            "by_jenis": by_jenis,
            "by_satker": by_satker,
            "per_unit": per_unit,
            "monitoring": monitoring,
        }
    finally:
        release_db(conn)

# --- AUTO NUMBER ---
@app.get("/api/auto-number/{doc_type}")
def get_auto_number(doc_type: str, tanggal: Optional[str] = None, unit: Optional[str] = None, user=Depends(get_current_user)):
    setting = query_one("SELECT value FROM settings WHERE type = 'format_nomor_surat'")
    if not setting or not setting.get("value"):
        raise HTTPException(status_code=404, detail="Format belum dikonfigurasi")
    formats = setting["value"]
    if doc_type not in formats:
        raise HTTPException(status_code=400, detail=f"Tipe '{doc_type}' tidak ditemukan")
    tgl = datetime.strptime(tanggal, "%Y-%m-%d").date() if tanggal else date.today()
    count_result = query_one("SELECT COUNT(*) as count FROM dumas WHERE EXTRACT(YEAR FROM created_at) = %s AND deleted_at IS NULL", (tgl.year,))
    next_number = (count_result["count"] if count_result else 0) + 1
    return {"number": auto_number(formats[doc_type], next_number, tgl, unit), "format": formats[doc_type]}

# --- DOCUMENT GENERATION ---
@app.get("/api/documents/{doc_type}/{dumas_id}")
def generate_document(doc_type: str, dumas_id: str, token: Optional[str] = None, user=Depends(get_current_user)):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    dumas = query_one("SELECT d.*, u.name as unit_name FROM dumas d LEFT JOIN users u ON d.unit_id = u.id WHERE d.id = %s::uuid", (dumas_id,))
    if not dumas:
        raise HTTPException(status_code=404, detail="Dumas tidak ditemukan")
    tl = query_one("SELECT * FROM tindak_lanjut WHERE dumas_id = %s::uuid ORDER BY created_at DESC LIMIT 1", (dumas_id,))
    
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    def add_header():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("KEPOLISIAN NEGARA REPUBLIK INDONESIA\nDAERAH JAWA BARAT\nBIDANG PROFESI DAN PENGAMANAN")
        run.bold = True; run.font.size = Pt(14)
    
    add_header()
    doc.add_paragraph()
    titles = {"uuk":"URAIAN UNSUR KEGIATAN (UUK)","sprin":"SURAT PERINTAH PENYELIDIKAN","sp2hp2":"SURAT PEMBERITAHUAN PERKEMBANGAN\nHASIL PENYELIDIKAN (SP2HP2)"}
    title = titles.get(doc_type, doc_type.upper())
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title); run.bold = True; run.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph(f"Nomor Dumas     : {dumas.get('no_dumas', '-')}")
    doc.add_paragraph(f"Tanggal Dumas   : {dumas.get('tgl_dumas', '-')}")
    doc.add_paragraph(f"Pelapor         : {dumas.get('pelapor', '-')}")
    doc.add_paragraph(f"Terlapor        : {dumas.get('terlapor', '-')}")
    doc.add_paragraph(f"Satker          : {dumas.get('satker', '-')}")
    doc.add_paragraph(f"Jenis Dumas     : {dumas.get('jenis_dumas', '-')}")
    doc.add_paragraph(f"Wujud Perbuatan : {dumas.get('wujud_perbuatan', '-')}")
    doc.add_paragraph(f"Keterangan      : {dumas.get('keterangan', '-')}")
    if tl:
        doc.add_paragraph(f"Hasil Lidik     : {tl.get('hasil_lidik', '-')}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"{doc_type}_{dumas.get('no_dumas', 'doc').replace('/', '_')}.docx"
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"})

# --- AUDIT LOG ---
@app.get("/api/audit-log")
def list_audit_log(page: int = 1, limit: int = 50, user=Depends(require_role("superadmin"))):
    offset = (page - 1) * limit
    return query_all(f"SELECT * FROM audit_log ORDER BY created_at DESC LIMIT {limit} OFFSET {offset}")

@app.get("/api/audit-log/count")
def count_audit_log(user=Depends(require_role("superadmin"))):
    result = query_one("SELECT COUNT(*) as total FROM audit_log")
    return {"total": result["total"] if result else 0}

# --- DUMAS PRINT DATA ---
@app.get("/api/dumas/{dumas_id}/print")
def get_dumas_print_data(dumas_id: str, user=Depends(get_current_user)):
    """Get complete dumas data for print view"""
    conn = get_db()
    try:
        cur = conn.cursor()
        cur.execute("SELECT d.*, u.name as unit_name FROM dumas d LEFT JOIN users u ON d.unit_id = u.id WHERE d.id = %s::uuid", (dumas_id,))
        cols = [d[0] for d in cur.description]
        dumas = serialize_row(cur.fetchone(), cols)
        
        cur.execute("SELECT * FROM tindak_lanjut WHERE dumas_id = %s::uuid AND deleted_at IS NULL ORDER BY created_at", (dumas_id,))
        tl_cols = [d[0] for d in cur.description]
        tindak_lanjut = [serialize_row(r, tl_cols) for r in cur.fetchall()]
        
        cur.execute("SELECT p.* FROM penyelesaian p WHERE p.dumas_id = %s::uuid ORDER BY created_at", (dumas_id,))
        p_cols = [d[0] for d in cur.description]
        penyelesaian_list = [serialize_row(r, p_cols) for r in cur.fetchall()]
        
        # Related dumas (children)
        cur.execute("SELECT id, no_dumas, pelapor, terlapor, status FROM dumas WHERE parent_dumas_id = %s::uuid AND deleted_at IS NULL", (dumas_id,))
        r_cols = [d[0] for d in cur.description]
        related = [serialize_row(r, r_cols) for r in cur.fetchall()]
        
        cur.close()
        return {"dumas": dumas, "tindak_lanjut": tindak_lanjut, "penyelesaian": penyelesaian_list, "related_dumas": related}
    finally:
        release_db(conn)
